import streamlit as st
import time
from io import BytesIO
from pypdf import PdfReader
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

from summarizer import (
    summarize_text,
    stream_summary
)

# =====================================
# PAGE CONFIG
# =====================================

st.set_page_config(
    page_title="AI Text Summarizer",
    page_icon="📄",
    layout="wide"
)

# =====================================
# MODELS
# =====================================

OLLAMA_MODELS = [
    "qwen2.5:1.5b",
    "gemma2:2b",
    "phi3:latest",
    "gemma3:4b",
    "mistral:latest",
    "llama3:8b"
]

OPENAI_MODELS = [
    "gpt-4o-mini",
    "gpt-4o"
]

# =====================================
# SESSION STATE
# =====================================

if "history" not in st.session_state:
    st.session_state.history = []

# =====================================
# TITLE
# =====================================

st.title("📄 AI Text Summarizer")
st.caption("Summarize text using Ollama or OpenAI")

# =====================================
# SIDEBAR
# =====================================

with st.sidebar:

    st.header("⚙️ Settings")

    provider = st.selectbox(
        "Provider",
        ["Ollama", "OpenAI"]
    )

    if provider == "Ollama":

        model_name = st.selectbox(
            "Choose Model",
            OLLAMA_MODELS,
            index=0
        )

        openai_api_key = None

    else:

        openai_api_key = st.text_input(
            "OpenAI API Key",
            type="password"
        )

        model_name = st.selectbox(
            "OpenAI Model",
            OPENAI_MODELS,
            index=0
        )

    st.markdown("---")

    st.header("📜 Summary History")

    if st.session_state.history:

        for i, item in enumerate(
            reversed(st.session_state.history),
            start=1
        ):

            st.write(
                f"{i}. {item['type']}"
            )

    else:
        st.caption("No summaries yet")

# =====================================
# SUMMARY TYPE
# =====================================

summary_type = st.selectbox(
    "Summary Type",
    [
        "Short",
        "Detailed",
        "Bullet Points"
    ]
)

# =====================================
# FILE UPLOAD
# =====================================

uploaded_file = st.file_uploader(
    "Upload TXT, PDF or DOCX",
    type=["txt", "pdf", "docx"]
)

uploaded_text = ""

if uploaded_file:

    suffix = uploaded_file.name.split(".")[-1].lower()

    try:

        if suffix == "txt":

            uploaded_text = (
                uploaded_file.read()
                .decode("utf-8")
            )

        elif suffix == "pdf":

            reader = PdfReader(
                uploaded_file
            )

            pages = []

            for page in reader.pages:

                pages.append(
                    page.extract_text() or ""
                )

            uploaded_text = "\n".join(
                pages
            )

        elif suffix == "docx":

            doc = Document(
                uploaded_file
            )

            uploaded_text = "\n".join(
                p.text
                for p in doc.paragraphs
            )

    except Exception as e:

        st.error(
            f"File Read Error: {e}"
        )

# =====================================
# TEXT AREA
# =====================================

text = st.text_area(
    "Paste your text here",
    value=uploaded_text,
    height=350,
    placeholder="Enter or paste text to summarize..."
)

# =====================================
# STATS
# =====================================

if text:

    word_count = len(
        text.split()
    )

    st.info(
        f"📊 Original Word Count: {word_count}"
    )

    if word_count > 1000:

        chunks = (
            word_count + 499
        ) // 500

        st.info(
            f"📦 Estimated Chunks: {chunks}"
        )

# =====================================
# GENERATE
# =====================================

if st.button(
    "Generate Summary",
    type="primary",
    use_container_width=True
):

    if not text.strip():

        st.warning(
            "Please enter text."
        )

        st.stop()

    if (
        provider == "OpenAI"
        and not openai_api_key
    ):

        st.warning(
            "Please enter OpenAI API Key."
        )

        st.stop()

    start_time = time.time()

    try:

        with st.spinner(
            "Generating summary..."
        ):

            summary = summarize_text(
                text=text,
                summary_type=summary_type,
                model_name=model_name,
                provider=provider,
                api_key=openai_api_key
            )

        elapsed = round(
            time.time() - start_time,
            2
        )

        st.success(
            f"✅ Summary Generated Successfully in {elapsed} seconds"
        )

        st.subheader(
            "📄 Summary"
        )

        placeholder = st.empty()

        summary = ""

        for token in stream_summary(
                text=text,
                summary_type=summary_type,
                model_name=model_name,
                provider=provider,
                api_key=openai_api_key
        ):
            summary += token

            placeholder.markdown(
                summary + "▌"
            )

        placeholder.markdown(
            summary
        )

        st.markdown(
            summary
        )

        # Save History

        st.session_state.history.append(
            {
                "type": summary_type,
                "provider": provider,
                "model": model_name,
                "summary": summary
            }
        )

        # =====================================
        # METRICS
        # =====================================

        original_words = len(
            text.split()
        )

        summary_words = len(
            summary.split()
        )

        reduction = round(
            (
                (
                    original_words
                    - summary_words
                )
                / original_words
            ) * 100,
            2
        )

        st.markdown("---")

        c1, c2, c3, c4 = st.columns(4)

        c1.metric(
            "Original Words",
            original_words
        )

        c2.metric(
            "Summary Words",
            summary_words
        )

        c3.metric(
            "Reduction %",
            f"{reduction}%"
        )

        c4.metric(
            "Time Taken",
            f"{elapsed}s"
        )

        # =====================================
        # TXT DOWNLOAD
        # =====================================

        st.download_button(
            "📥 Download TXT",
            summary,
            file_name="summary.txt",
            mime="text/plain"
        )

        # =====================================
        # PDF DOWNLOAD
        # =====================================

        pdf_buffer = BytesIO()

        pdf_doc = SimpleDocTemplate(
            pdf_buffer
        )

        styles = (
            getSampleStyleSheet()
        )

        story = [
            Paragraph(
                summary.replace(
                    "\n",
                    "<br/>"
                ),
                styles["BodyText"]
            )
        ]

        pdf_doc.build(
            story
        )

        pdf_buffer.seek(0)

        st.download_button(
            "📄 Download PDF",
            pdf_buffer,
            file_name="summary.pdf",
            mime="application/pdf"
        )

        # =====================================
        # DOCX DOWNLOAD
        # =====================================

        doc_buffer = BytesIO()

        doc = Document()

        doc.add_heading(
            "Summary",
            level=1
        )

        for line in summary.split("\n"):
            doc.add_paragraph(line)

        doc.save(doc_buffer)

        doc_buffer.seek(0)

        st.download_button(
            label="📝 Download DOCX",
            data=doc_buffer,
            file_name="summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:

        st.error(
            f"❌ Error: {str(e)}"
        )